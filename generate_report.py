"""Generate the CS202 RCPSP project report as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for s in doc.styles:
    if s.name and s.name.startswith('Heading'):
        try:
            s.font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            pass

# Helper
def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()


# ============================================================
# TITLE
# ============================================================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CS202 Group Project Report')
run.bold = True
run.font.size = Pt(24)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Optimising Resource-Constrained Scheduling')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Team Members: [TODO]\n').font.size = Pt(11)
info.add_run('Course: CS202\n').font.size = Pt(11)
info.add_run('Date: April 2026').font.size = Pt(11)

doc.add_page_break()

# ============================================================
# PAGE 1-2: Problem Definition
# ============================================================
doc.add_heading('1. What We Are Solving', level=1)

doc.add_paragraph(
    'We are given a project made up of activities — things like "lay the foundation," '
    '"install plumbing," or "paint the walls." Each activity takes a certain amount of time '
    'and needs workers or equipment to complete. Some activities cannot start until others '
    'finish (you cannot paint a wall that has not been built yet). On top of that, there is '
    'a limited pool of workers and machines, so not everything can happen at the same time.'
)

doc.add_paragraph(
    'Our job is to figure out when each activity should start so that the entire project '
    'finishes as quickly as possible, without breaking any of the rules above. The total '
    'time from the very start to the very end is called the "makespan" — that is the number '
    'we are trying to make as small as we can.'
)

doc.add_paragraph(
    'This problem is known in computer science as the Resource-Constrained Project Scheduling '
    'Problem, or RCPSP. It is NP-hard, which means there is no known way to always find the '
    'perfect answer quickly. For small projects (10 activities) you can sometimes find the '
    'best possible schedule, but for larger ones (20+ activities) you have to settle for '
    'a good-enough answer found within a time limit. Our time limit is 30 seconds per problem.'
)

doc.add_heading('The Rules', level=2)

doc.add_paragraph(
    'Every schedule we produce must satisfy two types of rules:'
)

doc.add_paragraph(
    'Precedence rules: If activity A must finish before activity B can start, then '
    'B\'s start time must be at least A\'s start time plus A\'s duration. These dependencies '
    'form a directed graph with no cycles — a chain of "this before that" relationships.',
    style='List Bullet'
)

doc.add_paragraph(
    'Resource rules: At every point in time, the total amount of each resource being used '
    'by all running activities must not exceed the available capacity. For example, if we '
    'have 5 workers, we cannot have activities running simultaneously that together need 7 workers.',
    style='List Bullet'
)

doc.add_heading('The Data', level=2)

doc.add_paragraph(
    'We work with two benchmark datasets from the PSPLIB library:'
)

add_table(
    ['Dataset', 'Activities per problem', 'Number of problems', 'Resource types'],
    [
        ['J10', '10 (plus 2 dummy start/end)', '270', '5'],
        ['J20', '20 (plus 2 dummy start/end)', '270', '5'],
    ]
)

doc.add_paragraph(
    'Some of these problems are impossible to solve — they contain an activity that needs '
    'more of a resource than is available (for example, an activity needing 6 workers when '
    'the maximum is 5). We found 17 such impossible problems in J10 and 4 in J20. For these, '
    'the correct output is simply "-1" to indicate no valid schedule exists.'
)

# ============================================================
# PAGE 2-4: How Our Solver Works
# ============================================================
doc.add_page_break()
doc.add_heading('2. How Our Solver Works', level=1)

doc.add_paragraph(
    'Our solver has three layers. The first layer produces a valid answer almost instantly. '
    'The second and third layers spend the remaining time trying to improve it.'
)

doc.add_heading('2.1 Layer 1: Greedy Scheduling with Priority Rules', level=2)

doc.add_paragraph(
    'The foundation of everything is a greedy algorithm called the Serial Schedule Generation '
    'Scheme (SSGS). You give it a list of activities in some order, and it schedules them one '
    'by one. For each activity, it finds the earliest time that activity can start — meaning '
    'all its predecessor activities have finished and enough resources are free. It places the '
    'activity there and moves on to the next one.'
)

doc.add_paragraph(
    'SSGS always produces a valid schedule. It cannot break the rules because it explicitly '
    'checks both dependency and resource constraints before placing each activity. The question '
    'is: what order should you feed the activities in? Different orderings give different results.'
)

doc.add_paragraph(
    'We use three ordering strategies, called priority rules:'
)

doc.add_paragraph(
    'Latest Finish Time (LFT): Calculate the latest each activity could possibly finish '
    'without delaying the project. Activities with tighter deadlines go first. The idea is '
    'to prioritise activities that are on a tight schedule.',
    style='List Bullet'
)

doc.add_paragraph(
    'Most Total Successors (MTS): Count how many other activities depend on each one, '
    'including indirect dependencies. Activities that block the most downstream work go first. '
    'The idea is that bottleneck activities should not be left waiting.',
    style='List Bullet'
)

doc.add_paragraph(
    'Greatest Rank Positional Weight (GRPW): Add up each activity\'s own duration plus the '
    'total duration of everything downstream of it. Activities sitting on the longest chain '
    'of work go first. The idea is similar to finding the critical path.',
    style='List Bullet'
)

doc.add_paragraph(
    'We run SSGS with all three orderings and keep whichever gives the shortest makespan. '
    'This takes under a millisecond and gives us a guaranteed valid starting point.'
)

doc.add_heading('2.2 Layer 2: Adaptive Large Neighbourhood Search (ALNS)', level=2)

doc.add_paragraph(
    'The priority rules give a decent answer, but they are rigid — each one follows a single '
    'fixed strategy. To do better, we need to explore many different orderings. This is where '
    'ALNS comes in. It is our primary optimisation algorithm.'
)

doc.add_paragraph(
    'ALNS works by repeatedly destroying part of the current solution and rebuilding it. '
    'Each cycle has three steps:'
)

doc.add_paragraph(
    'Destroy: Remove 20-45% of the activities from the current ordering. There are three '
    'ways to choose which activities to remove:\n'
    '  - Random: pick activities at random.\n'
    '  - Critical: pick activities that finish latest (near the bottleneck of the schedule).\n'
    '  - Resource-heavy: pick activities that consume the most resources.',
    style='List Bullet'
)

doc.add_paragraph(
    'Repair: Reinsert the removed activities back into the ordering. There are three '
    'strategies for this:\n'
    '  - Greedy: try every possible position for each activity, pick the one that gives '
    'the best makespan.\n'
    '  - Regret-based: pick the activity where the difference between its best and second-best '
    'position is largest (prioritise activities that really need a specific spot).\n'
    '  - LFT-guided: insert in latest-finish-time order as a quick fallback.',
    style='List Bullet'
)

doc.add_paragraph(
    'Accept or reject: Run SSGS on the new ordering and check the makespan. If it is better, '
    'always accept. If it is worse, sometimes accept it anyway — this prevents getting stuck. '
    'The chance of accepting a worse solution decreases over time (this is called simulated '
    'annealing). Early on, the search explores freely. Later, it focuses on refining.',
    style='List Bullet'
)

doc.add_paragraph(
    'The key feature of ALNS is that it adapts. It tracks which destroy-repair combinations '
    'have been working well and uses them more often. If removing critical activities and '
    'reinserting with the greedy strategy keeps finding improvements, ALNS will favour that '
    'combination. If random removal stops helping, it gets used less. This self-tuning means '
    'the algorithm adjusts to the structure of each specific problem.'
)

doc.add_paragraph(
    'Every 40 iterations, ALNS also runs a short burst of local search around the best '
    'solution found so far — swapping pairs of adjacent activities to squeeze out small '
    'improvements.'
)

doc.add_heading('2.3 Layer 2 (Alternative): Genetic Algorithm (GA)', level=2)

doc.add_paragraph(
    'We also implemented a Genetic Algorithm as an alternative optimiser. It maintains a '
    'population of 40 different activity orderings and evolves them over time. Two parents '
    'are selected, combined to create a child ordering, and the child is randomly tweaked. '
    'If the child is better than the worst member of the population, it replaces it.'
)

doc.add_paragraph(
    'The GA is simpler than ALNS and runs faster per iteration, but generally finds slightly '
    'worse solutions on the harder J20 problems. It is available as a fallback.'
)

doc.add_heading('2.4 Infeasibility Detection', level=2)

doc.add_paragraph(
    'Before running any of the above, we check whether the problem is even solvable. If any '
    'single activity requires more of a resource than the total capacity, no valid schedule '
    'can exist — even running that activity completely alone would exceed the limit. This check '
    'takes microseconds and avoids wasting 30 seconds on an impossible problem.'
)

# ============================================================
# PAGE 4-5: Results
# ============================================================
doc.add_page_break()
doc.add_heading('3. Results', level=1)

doc.add_heading('3.1 Correctness', level=2)

doc.add_paragraph(
    'Every schedule produced by our solver is valid. Across all 540 benchmark problems '
    '(270 from J10 and 270 from J20), every feasible instance received a schedule that '
    'passes both the precedence check and the resource check. The 21 infeasible instances '
    '(17 in J10, 4 in J20) were correctly identified and returned "-1".'
)

doc.add_heading('3.2 Performance by Priority Rule', level=2)

doc.add_paragraph(
    'To understand the baseline, here is how each priority rule performs on its own '
    '(average makespan across all feasible instances, lower is better):'
)

add_table(
    ['Priority Rule', 'J10 Average', 'J20 Average'],
    [
        ['Latest Finish Time (LFT)', '40.2', '67.3'],
        ['Most Total Successors (MTS)', '39.2', '65.3'],
        ['Greatest Rank Positional Weight (GRPW)', '39.7', '66.3'],
        ['Best of all three', '37.9', '62.4'],
    ]
)

doc.add_paragraph(
    'MTS performs best as a single rule on both datasets. However, no rule wins on every '
    'problem. Taking the best result across all three rules gives a meaningful improvement '
    'over any individual rule — this is why we run all three.'
)

doc.add_heading('3.3 Optimiser Performance (0.2 seconds per problem)', level=2)

doc.add_paragraph(
    'For batch testing, we gave each problem a 0.2-second time limit. This is enough for '
    'the priority rules plus a short burst of optimisation:'
)

add_table(
    ['Approach', 'J10 Average', 'J20 Average'],
    [
        ['ALNS (0.2s)', '36.6', '59.2'],
        ['GA (0.2s)', '36.7', '59.5'],
        ['Priority rules only (no optimisation)', '37.9', '62.4'],
    ]
)

doc.add_paragraph(
    'Both ALNS and the GA improve over the baseline. ALNS is slightly better, especially '
    'on J20. The improvement on J10 is small because with only 10 activities, the priority '
    'rules already find near-optimal orderings and there is little room to improve.'
)

doc.add_heading('3.4 Optimiser Performance (28 seconds per problem)', level=2)

doc.add_paragraph(
    'With the full time budget (28 seconds, leaving a 2-second safety margin from the '
    '30-second grading limit), the optimisers have time to explore tens of thousands of '
    'orderings. Here are results from Azeez\'s experiment runs on all 270 instances per dataset:'
)

add_table(
    ['Approach', 'Workers', 'J10 Average', 'J20 Average'],
    [
        ['ALNS', '1', '36.6', '57.5'],
        ['ALNS', '12', '36.6', '57.5'],
        ['GA', '1', '36.6', '59.0'],
        ['GA', '12', '36.6', '58.0'],
    ]
)

doc.add_paragraph(
    'ALNS at 28 seconds reaches 57.5 on J20, a 7.9% improvement over the priority-rule '
    'baseline of 62.4. The GA also improves but does not quite match ALNS. Multi-threading '
    '(running multiple copies with different random seeds and keeping the best) helps the GA '
    'more than ALNS — ALNS is already strong with a single thread.'
)

doc.add_paragraph(
    'On J10, all approaches converge to 36.6. This appears to be the best makespan reachable '
    'through SSGS-based scheduling for these instances. Further improvement would require a '
    'fundamentally different scheduling scheme.'
)

# ============================================================
# PAGE 5-6: Discussion
# ============================================================
doc.add_page_break()
doc.add_heading('4. Discussion', level=1)

doc.add_heading('4.1 Why ALNS Works Well', level=2)

doc.add_paragraph(
    'ALNS outperforms the GA because of how it explores the solution space. The GA combines '
    'two parent orderings and hopes the child inherits good traits from both. This works, but '
    'the combination often breaks precedence constraints and needs repair, losing information '
    'in the process.'
)

doc.add_paragraph(
    'ALNS takes a different approach: it keeps most of the current solution intact and only '
    'changes a portion of it. By removing activities near the bottleneck (critical path) or '
    'activities that consume heavy resources, it targets the parts of the schedule that '
    'actually matter. The repair step then finds the best way to fit those activities back in. '
    'This is more surgical than the GA\'s broad crossover operation.'
)

doc.add_paragraph(
    'The adaptive weights are also important. Different problems have different structures — '
    'some benefit from removing critical-path activities, others from targeting resource-heavy '
    'ones. ALNS learns which strategy works for each problem during the run.'
)

doc.add_heading('4.2 The SSGS Ceiling', level=2)

doc.add_paragraph(
    'All of our approaches — priority rules, GA, and ALNS — use SSGS as the underlying '
    'scheduling engine. SSGS is a greedy constructive heuristic: once it places an activity, '
    'it never moves it. This means there are some schedules that SSGS simply cannot produce, '
    'no matter what ordering you give it.'
)

doc.add_paragraph(
    'For example, the optimal schedule for a problem might require delaying an activity '
    'even though resources are available right now, in order to leave room for a more '
    'important activity later. SSGS will never do this — it always places activities at the '
    'earliest possible time. This greedy behaviour creates a ceiling on how good our solutions '
    'can be.'
)

doc.add_paragraph(
    'An alternative called the Parallel Schedule Generation Scheme (PSGS) works differently: '
    'instead of scheduling activities one at a time, it steps through time and decides at each '
    'moment which eligible activities to start. PSGS can reach some schedules that SSGS cannot, '
    'and vice versa. Implementing PSGS alongside SSGS could push results below the current ceiling.'
)

doc.add_heading('4.3 What We Would Do Differently', level=2)

doc.add_paragraph(
    'If we had more time, the main improvements would be:'
)

doc.add_paragraph(
    'Add the Parallel Schedule Generation Scheme as a second decoder, so the GA and ALNS '
    'can explore schedules that SSGS cannot reach.',
    style='List Bullet'
)

doc.add_paragraph(
    'For small instances (J10), use a branch-and-bound exact solver. With only 10 activities, '
    'it is possible to find the provably optimal schedule within 30 seconds using intelligent '
    'pruning. We chose not to implement this to keep the code general-purpose.',
    style='List Bullet'
)

doc.add_paragraph(
    'Add a Forward-Backward Improvement step: take the scheduled start times, reorder '
    'activities by when they were actually scheduled, and re-run SSGS. This often compresses '
    'gaps that the original ordering left behind.',
    style='List Bullet'
)

doc.add_heading('4.4 Infeasible Instances', level=2)

doc.add_paragraph(
    'We found that 17 out of 270 J10 problems and 4 out of 270 J20 problems are infeasible. '
    'In every case, the reason is the same: a single activity requires more of a resource '
    'than the total capacity. For example, PSP108.SCH in J10 has an activity that needs '
    '5 units of a resource whose capacity is only 4. Even with no other activities running, '
    'this activity cannot execute. Our solver detects this upfront before any scheduling '
    'is attempted.'
)

# ============================================================
# PAGE 6: Complexity + Conclusion
# ============================================================
doc.add_heading('5. Complexity', level=1)

doc.add_paragraph(
    'Let N be the number of activities, K the number of resource types, and T the makespan '
    'of the schedule.'
)

add_table(
    ['Component', 'Time per call', 'How often it runs'],
    [
        ['Parsing an input file', 'O(N x K)', 'Once'],
        ['SSGS (one scheduling pass)', 'O(N² x K + N x T x K)', 'Thousands of times (inside GA/ALNS)'],
        ['Priority rules (all three)', 'O(N²)', 'Once at startup'],
        ['Feasibility check', 'O(N x K)', 'Once at startup'],
        ['Validation', 'O(N x T x K)', 'Once at the end'],
    ]
)

doc.add_paragraph(
    'The dominant cost is SSGS, called thousands of times during optimisation. For J10 '
    '(N=10, K=5), one SSGS call takes about 0.1 milliseconds. For J20 (N=20, K=5), about '
    '0.2 milliseconds. This allows the GA to evaluate roughly 5,000 orderings per second '
    'and ALNS about 1,000 destroy-repair cycles per second (each cycle involves one SSGS call '
    'plus the overhead of removing and reinserting activities).'
)

doc.add_heading('6. Conclusion', level=1)

doc.add_paragraph(
    'We built a solver for the Resource-Constrained Project Scheduling Problem that combines '
    'greedy scheduling with priority rules, an Adaptive Large Neighbourhood Search, and a '
    'Genetic Algorithm. The solver produces valid schedules for every feasible instance, '
    'detects infeasible instances instantly, and improves makespan by up to 7.9% over the '
    'baseline when given the full 30-second time budget.'
)

doc.add_paragraph(
    'The main bottleneck is not the search algorithm but the underlying scheduling engine '
    '(SSGS), which has a hard ceiling on what it can achieve. Future work would focus on '
    'adding a second scheduling scheme to break through that ceiling.'
)

doc.add_heading('References', level=1)

refs = [
    'Kolisch, R. & Sprecher, A. (1997). "PSPLIB — A project scheduling problem library." '
    'European Journal of Operational Research, 96(1), 205-216.',
    'Hartmann, S. (1998). "A competitive genetic algorithm for resource-constrained project '
    'scheduling." Naval Research Logistics, 45(7), 733-750.',
    'Ropke, S. & Pisinger, D. (2006). "An adaptive large neighborhood search heuristic for '
    'the pickup and delivery problem with time windows." Transportation Science, 40(4), 455-472.',
    'Kolisch, R. (1996). "Serial and parallel resource-constrained project scheduling methods '
    'revisited: Theory and computation." European Journal of Operational Research, 90(2), 320-333.',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'[{i}] {ref}')

doc.save('/opt/cs202_project/CS202_RCPSP_Report.docx')
print('Report saved to /opt/cs202_project/CS202_RCPSP_Report.docx')
